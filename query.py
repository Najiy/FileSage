import sys, os
from langchain_community.vectorstores import FAISS
from langchain_openai import OpenAIEmbeddings
from docx import Document
from docx.shared import Pt, Inches
from docx2pdf import convert
import markdown2
from bs4 import BeautifulSoup
from dotenv import load_dotenv

load_dotenv()

DB_DIR = 'db/'
REPORTS_DIR = 'reports/'
os.makedirs(REPORTS_DIR, exist_ok=True)

def query_db(question):
    embeddings = OpenAIEmbeddings(model="text-embedding-3-small")
    db = FAISS.load_local(DB_DIR, embeddings, allow_dangerous_deserialization=True)
    return db.similarity_search(question, k=6)

def render_markdown_to_docx(doc, markdown_text, base_path):
    html = markdown2.markdown(markdown_text)
    soup = BeautifulSoup(html, "html.parser")
    for element in soup.contents:
        if element.name == 'h1':
            doc.add_heading(element.get_text(), level=1)
        elif element.name == 'h2':
            doc.add_heading(element.get_text(), level=2)
        elif element.name == 'h3':
            doc.add_heading(element.get_text(), level=3)
        elif element.name == 'ul':
            for li in element.find_all('li'):
                doc.add_paragraph(li.get_text(), style='List Bullet')
        elif element.name == 'ol':
            for li in element.find_all('li'):
                doc.add_paragraph(li.get_text(), style='List Number')
        elif element.name == 'img':
            src = element.get('src')
            img_path = os.path.join(base_path, src)
            if os.path.exists(img_path):
                doc.add_picture(img_path, width=Inches(4))
                doc.add_paragraph(f"(Image: {src})").italic = True
        elif element.name == 'p':
            doc.add_paragraph(element.get_text())
        elif element.name == 'pre':
            doc.add_paragraph(element.get_text(), style='Intense Quote')
        else:
            if hasattr(element, 'get_text'):
                doc.add_paragraph(element.get_text())

def create_docx(query, docs, docx_filename):
    doc = Document()
    doc.add_heading(f'Query: {query}', level=1)

    for idx, d in enumerate(docs, 1):
        source = d.metadata['source']
        doc_type = d.metadata.get('type', 'text')
        ext = os.path.splitext(source)[1].lower()
        base_path = os.path.dirname(source)

        doc.add_heading(f'Result #{idx}', level=2)

        if doc_type == 'image':
            doc.add_paragraph("Extracted Text from Image:", style='Intense Quote')
            doc.add_paragraph(d.page_content)
            if os.path.exists(source):
                doc.add_picture(source, width=Inches(4))
        elif ext == '.md':
            render_markdown_to_docx(doc, d.page_content, base_path)
        else:
            doc.add_paragraph(d.page_content)

        p = doc.add_paragraph()
        p.add_run("Source: ").italic = True
        p.add_run(source).bold = True

        doc.add_paragraph("\n")

    doc.save(docx_filename)
    print(f"✅ DOCX saved: {docx_filename}")

def create_pdf(docx_filename, pdf_filename):
    convert(docx_filename, pdf_filename)
    print(f"✅ PDF saved: {pdf_filename}")

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python query.py \"Your question here\"")
        sys.exit(1)

    query = sys.argv[1]
    results = query_db(query)

    for r in results:
        print(f"\nSource: {r.metadata['source']}\n---\n{r.page_content}\n")

    base = os.path.join(REPORTS_DIR, query.replace(" ", "_"))
    docx_path = base + ".docx"
    pdf_path = base + ".pdf"

    create_docx(query, results, docx_path)
    create_pdf(docx_path, pdf_path)
